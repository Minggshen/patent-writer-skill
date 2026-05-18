"""
专利初稿Word文档生成脚本

将结构化专利内容输出为格式规范的.docx技术交底书文档。
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr in ['sz', 'val', 'color', 'space']:
                if attr in edge_data:
                    element.set(qn(f'w:{attr}'), str(edge_data[attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def create_patent_document(patent_data, output_path):
    """
    生成专利初稿Word文档

    patent_data = {
        'title': '专利名称',
        'inventors': '发明人',
        'applicant': '申请人',
        'address': '地址',
        'field': '应用领域',
        'problem': '要解决的技术问题',
        'prior_art': '已有技术的不足',
        'technical_content': [
            {'subtitle': '总体方案', 'content': '...', 'figures': [('图1', '描述')]},
            {'subtitle': '具体技术内容1', 'content': '...', 'figures': [...]},
            ...
        ],
        'figures_list': ['图1 描述', '图2 描述', ...],
        'effects': '有益效果',
    }
    """
    doc = Document()

    # Document default style setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Paragraph format
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    # Heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        hf = heading_style.font
        hf.name = 'Times New Roman'
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        hf.color.rgb = RGBColor(0, 0, 0)
        if i == 1:
            hf.size = Pt(16)
        elif i == 2:
            hf.size = Pt(14)
        else:
            hf.size = Pt(13)

    # ===== Title =====
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(patent_data.get('title', ''))
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = 'Times New Roman'
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # ===== Inventor/Applicant/Address =====
    for key, label in [('inventors', '发明人：'), ('applicant', '申请（专利权）人：'), ('address', '地址：')]:
        if patent_data.get(key):
            p = doc.add_paragraph()
            p.add_run(label).bold = True
            p.add_run(patent_data[key])

    # ===== Section: Application Field =====
    doc.add_heading('本专利的应用领域', level=2)
    doc.add_paragraph(patent_data.get('field', ''))

    # ===== Section: Technical Problem =====
    doc.add_heading('本专利的任务是什么，或要解决的技术问题是什么？', level=2)
    doc.add_paragraph(patent_data.get('problem', ''))

    # ===== Section: Prior Art =====
    doc.add_heading('已有技术/产品的不足', level=2)
    doc.add_paragraph(patent_data.get('prior_art', ''))

    # ===== Section: Technical Content =====
    doc.add_heading('本专利的内容', level=2)

    for section in patent_data.get('technical_content', []):
        subtitle = section.get('subtitle', '')
        content = section.get('content', '')
        figures = section.get('figures', [])

        if subtitle:
            doc.add_heading(subtitle, level=3)

        # Split content by paragraphs and add
        for para_text in content.split('\n'):
            para_text = para_text.strip()
            if para_text:
                doc.add_paragraph(para_text)

        # Add figure placeholders
        for fig_label, fig_desc in figures:
            fig_para = doc.add_paragraph()
            fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_run = fig_para.add_run(f'[{fig_label}] {fig_desc}')
            fig_run.italic = True
            fig_run.font.size = Pt(10)
            fig_run.font.color.rgb = RGBColor(128, 128, 128)

    # ===== Section: Figures =====
    doc.add_heading('附图或图表', level=2)
    for fig in patent_data.get('figures_list', []):
        doc.add_paragraph(fig, style='List Number')

    # ===== Section: Effects =====
    doc.add_heading('本专利的效果', level=2)
    doc.add_paragraph(patent_data.get('effects', ''))

    # ===== Save =====
    doc.save(output_path)
    print(f'专利文档已保存至: {output_path}')
    return output_path


if __name__ == '__main__':
    # Example usage
    data = {
        'title': '一种基于示例方法的电力系统优化调度方法',
        'inventors': '',
        'applicant': '',
        'address': '',
        'field': '本发明涉及电力系统优化调度领域，尤其涉及一种基于示例方法的电力系统优化调度方法。',
        'problem': '此处填写技术问题...',
        'prior_art': '此处填写已有技术的不足...',
        'technical_content': [
            {
                'subtitle': '1. 总体方案',
                'content': '本发明提出一种...',
                'figures': [('图1', '整体技术路线图')]
            },
        ],
        'figures_list': ['图1 整体技术路线图'],
        'effects': '此处填写有益效果...',
    }
    create_patent_document(data, '示例专利初稿.docx')
